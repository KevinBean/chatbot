# -*- coding=utf-8 -*-
'''
生成一份docx格式的说明书。
使用md格式书写文档，使用pypandoc输出docx文件。
'''
import pypandoc as pydoc
from caculate import *
from attract import *
import pickle
'''
from Tkinter import *

root = Tk()
root.title(u'说明书自动生成')

# 创建几个frame作为容器
frame_toplevel1 = Frame(width=380, height=270, bg='red')

# 创建需要的几个元素

    # 2. 创建主窗体
    root = Tk()
    root.title(unicode('对话窗口', 'utf-8'))

    # 创建几个frame作为容器
    frame_top = Frame(width=380, height=270, bg='red')
    # frame_left_center  = Frame(width=380, height=100, bg='yellow')
    frame_left_bottom = Frame(width=250, height=50)
    frame_center_bottom = Frame(width = 70,height =50)
    frame_right_bottom = Frame(width=60, height=50)
    # frame_right     = Frame(width=170, height=400, bg='white')

    ##创建需要的几个元素
    # text_msglist作为对话信息展示区域，在sendmessage函数中处理
    text_msglist = ScrolledText(frame_top,wrap=WORD)  #加滑动条的文本区，wrap自动换行
    msg_lable = Label(frame_left_bottom,text='请输入')
    # text_msg作为信息输入框按钮，其中信息被read_msg()函数读取
    text_msg = Entry(frame_left_bottom,width=24)
    select_lable =Label(frame_center_bottom,text='请选择')
    text_select = Listbox(frame_center_bottom,width=20,height = 2)
    button_lable = Label(frame_right_bottom,text='按我')
    # button_sendmsg作为发送按钮，触发sendmessage函数
    button_sendmsg = Button(frame_right_bottom, text=unicode('发送', 'utf-8'), command=sendmessage)
    # 不显示的标签
    text_respond = Label(frame_top)

    # 创建一个绿色的tag
    text_msglist.tag_config('green', foreground='#008B00')
    # 3. 显示欢迎信息
    insert_msg('Robot',u'您好,请问有什么可以帮您:)')

    # 使用grid设置各个容器位置
    frame_top.grid(row=0, column=0, padx=2, pady=5, columnspan=3,sticky=N)
    # frame_left_center.grid(row=1, column=0, padx=2, pady=5)
    frame_left_bottom.grid(row=1, column=0,sticky=W)
    frame_center_bottom.grid(row=1,column=1,sticky=E)
    frame_right_bottom.grid(row=1, column=2, padx=1, pady=2,sticky=E)
    frame_top.grid_propagate(0)
    frame_left_bottom.grid_propagate(0)
    frame_right_bottom.grid_propagate(0)

    # 把元素填充进frame
    text_msglist.pack(expand=True, fill='both') #文本区加滑动条
    msg_lable.grid()
    text_msg.grid()
    select_lable.grid()
    text_select.grid()
    button_lable.grid(sticky=W)
    button_sendmsg.grid(sticky=E)


    # 4. 主事件(窗体)循环，等待事件触发
    root.mainloop()


'''


'''
from docx import Document #使用python-docx包

document = Document()

# 电缆路径
document.add_heading(u'电缆路径',level=1)
p = document.add_paragraph(u'本工程新建')
p.add_run(u'双回110kV电缆')

# 电缆电气部分
'''


pkl = '0371'

'''
#生成测试文件
sms = shuomingshu(n=1)
cs = sms.dianlan()
print(cs.Cable['电缆型号'])
# 保存数据以备调用
pklfile = 'dict/' + pkl +'.pkl'
output = open(pklfile, 'wb')
pickle.dump(cs,output)
output.close()
'''



#读取提资文件中系统概况
if os.name == 'nt':
    filename = u'D:\Personal\我的文档\GitHub\chatbot\doc\X9348K-X-02 互提资料单 送电.txt' #系统提资文件路径
    dictpath = u'D:\Personal\我的文档\GitHub\chatbot\dict\dict.txt'  # 训练数据越多越准确
else:
    filename = u'doc/X9348K-X-02 互提资料单 送电.txt'  # 系统提资文件路径
    filename = u'doc/北河系统提资（送电）.txt'  # 系统提资文件路径
    dictpath = u'dict/dict.txt'  # 训练数据越多越准确
dictionary = corpora.Dictionary.load_from_text(dictpath)

#读取保存的测试数据
readcs = open(pklfile,'r')
cs = pickle.load(readcs)
readcs.close()
#生成测试文件
sms = shuomingshu(n=1)
cs = sms.dianlan(filename,dictpath)


sample_xitonggaikuang = u'将村齐线“T”接南山的110kV线路破口接入永定变电站，形成永定～南山以及永定～规划园博园的110kV线路。本期工程完成后，南山站电源将由永定站提供；园博园站电源也由永定站主供，吕村站作为备用电源。'
xtgk = findinFile(filename, dictionary, sample_xitonggaikuang)
print xtgk
#生成系统概况段落
xitonggaikuang = ''
xitonggaikuang += u'## 系统概况'
xitonggaikuang += '\n'
xitonggaikuang += xtgk
xitonggaikuang += '\n'
xitonggaikuang += '\n'

lujing = ''
lujing += u'## 电缆路径'
lujing += '\n'
lujing += u'本工程新建'
lujing += u'双回110kV电缆'
lujing += u'自' + cs.startPoint['名称'].decode('utf-8') # 替换成起点名称
lujing += '' #替换成路径描述 电缆路径描述是可以根据地图自动生成的
lujing += u'至' + cs.endPoint['名称'].decode('utf-8')  + u'。' # 替换成终点名称
lujing += '\n'
lujing += '\n'

dianqibufen = ''
dianqibufen += u'## 电缆电气部分'
dianqibufen += '\n'
dianqibufen += '\n'
dianqibufen += u'### 电缆及附件选型'
dianqibufen += '\n'
dianqibufen += '\n'
dianqibufen += u'#### 电缆选型'
dianqibufen += '\n'
dianqibufen += '\n'
dianqibufen += u'根据系统要求'
dianqibufen += cs.info[u'系统载流量提资'] #替换成系统专业的提取值
dianqibufen += u'新建电缆'
dianqibufen += u'全线敷设在电力隧道及电缆夹层中，' #替换成敷设环境
dianqibufen += u'在环境温度40℃、线芯运行温度90℃、品字形接触排列的条件下，' #替换成敷设环境
dianqibufen += u'参考现有厂家资料，'
dianqibufen += u'铜芯、交联' + str(cs.Cable['截面']).decode('utf-8') + u'平方毫米截面电缆'
dianqibufen += u'，载流量不小于' + str(cs.Cable['最大载流量']).decode('utf-8') + u'A，' #替换成电缆型号及载流量
dianqibufen += u'满足系统载流量要求，'
dianqibufen += u'故本工程选用'
dianqibufen += cs.Cable['电缆型号'].decode('utf-8') #替换成电缆型号
dianqibufen += u'电缆。'
dianqibufen += '\n'
dianqibufen += '\n'
dianqibufen += u'#### 附件选型'
dianqibufen += '\n'
dianqibufen += '\n'
dianqibufen += u'全线敷设在电力隧道及电缆夹层中，' #替换成敷设环境
dianqibufen += u'全线敷设在电力隧道及电缆夹层中，' #替换成敷设环境
dianqibufen += u'全线敷设在电力隧道及电缆夹层中，' #替换成敷设环境
dianqibufen += u'全线敷设在电力隧道及电缆夹层中，' #替换成敷设环境
dianqibufen += '\n'
dianqibufen += '\n'
dianqibufen += u'#### 附件选型'
dianqibufen += '\n'
print cs.MaterialType['户外终端']
dianqibufen += u'户外终端' + cs.MaterialType['户外终端']
dianqibufen += '\n'
dianqibufen += '\n'


gongzuoliang = ''
gongzuoliang += u'主要工作量：'
gongzuoliang += '\n'
gongzuoliang += '\n'
gongzuoliang += '|' + u'敷设：' + '|' + '|' + '|'
gongzuoliang += '\n'
gongzuoliang += '|' + '------' + '|' + '------' + '|' + '------' + '|'
gongzuoliang += '\n'
gongzuoliang += '|' + '12' + '|' + '34' + '|' + '56' + '|'
gongzuoliang += '\n'
gongzuoliang += '|' + '65' + '|' + '43' + '|' + '21' + '|'
gongzuoliang += '\n'
gongzuoliang += '|' + u'安装：' + '|' + '|' + '|'
gongzuoliang += '\n'
gongzuoliang += '|' + u'户外终端' + '|' + cs.MaterialType['户外终端'] + '|' + str(cs.Material['户外终端']) + '|'
gongzuoliang += '\n'

gongzuoliang += '\n'
gongzuoliang += '\n'

jiedifangshi = ''
jiedifangshi += u'## 电缆接地方式'
jiedifangshi += '\n'
jiedifangshi += cs.Cable['接地方式']



output = xitonggaikuang + lujing + dianqibufen +gongzuoliang + jiedifangshi

'''
output = open ('test.md', 'w')
output.write(dianlanlujing)
output.close()
'''

pydoc.convert_text(output,to='docx',format = 'md',outputfile='doc/test.docx')