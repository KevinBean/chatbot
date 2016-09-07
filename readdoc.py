# -*- coding=utf-8 -*-
# import pandas
# 直接改成pandoc转换更便捷
#
from docx import Document #使用python-docx包
import jieba # 利用jieba进行中文分词
import nltk  # 利用nltk进行其他处理
import pypandoc
import os
if os.name == 'nt':
    from win32com import client as wc


def docx2txt(filename):
    newfilename = filename.replace(u'docx', u'txt')
    if os.name == 'nt':
        print 'nt'
        word = wc.Dispatch('Word.Application')
        doc = word.Documents.Open(filename)
        doc.SaveAs(newfilename, 4)
        doc.Close()
        word.Quit() #另存为txt文件，编码为gbk
        input_file = open(newfilename, 'r')
        gbktxt = input_file.read()
        utftxt = gbktxt.decode('gbk').encode('utf-8') #读取txt文件，将gbk转换成utf-8
        input_file.close()
        output_file = open(newfilename, 'w')
        output_file.write(utftxt) #保存utf-8文本
        output_file.close()

    else:
        '''
        # 从word（docx格式）中提取text，保存为txt
        document = Document(filename)
        docText = '\n\n'.join([
                                  paragraph.text.encode('utf-8') for paragraph in document.paragraphs
                                  ])
        print docText
        # 保存文件
        # document.save('doc/new-SL351C-A11-01.doc')

        output_file = open(newfilename, 'w')
        output_file.write(docText)
        output_file.close()
        '''
        #使用pandoc进行转换
        pypandoc.convert_file(filename,'markdown','docx',outputfile=newfilename)
        print newfilename


def doc2txt(filename):
    newfilename = filename.replace(u'docx', u'txt')
    if os.name == 'nt':
        print 'nt'
        word = wc.Dispatch('Word.Application')
        doc = word.Documents.Open(filename)
        doc.SaveAs(newfilename, 4)
        doc.Close()
        word.Quit() #另存为txt文件，编码为gbk
        input_file = open(newfilename, 'r')
        gbktxt = input_file.read()
        utftxt = gbktxt.decode('gbk').encode('utf-8') #读取txt文件，将gbk转换成utf-8
        input_file.close()
        output_file = open(newfilename, 'w')
        output_file.write(utftxt) #保存utf-8文本
        output_file.close()

    else:
        # 从word（docx格式）中提取text，保存为txt
        document = Document(filename)
        docText = '\n\n'.join([
                                  paragraph.text.encode('utf-8') for paragraph in document.paragraphs
                                  ])
        print docText
        # 保存文件
        # document.save('doc/new-SL351C-A11-01.doc')

        output_file = open(newfilename, 'w')
        output_file.write(docText)
        output_file.close()
    '''word = wc.Dispatch('Word.Application')
    #filepath = os.path.abspath('.').decode('gbk').encode('utf-8') #获取文件绝对路径
    #print filepath,type(filepath),type(filename)
    #filename = os.path.join(filepath.decode('utf-8') + filename).replace('/','\\').decode('gbk').encode('utf-8')
    print filename
    doc = word.Documents.Open(filename) # 所有Paragraph即段落对象，都是通过Paragraph.Range.Text来访问它的文字的
    docText = '\n\n'.join([
                              paragraph.Range.Text.encode('utf-8') for paragraph in doc.paragraphs
                              ])
    word.Quit()
    newfilename = filename.replace(u'.doc', u'.txt')
    print newfilename
    output_file = open(newfilename, 'w')
    output_file.write(docText)
    output_file.close()'''

if __name__ == '__main__':
    #filename = ur'D:\Personal\我的文档\GitHub\chatbot\doc\X9384K-AB-01.doc' # u'doc/管廊缆线敷设技术条件.docx'
    if os.name == 'nt':
        filename = u'D:\Personal\我的文档\GitHub\chatbot\doc\X9348K-X-02 互提资料单 送电.docx'
        filename = os.path.normpath(filename)
    else:
        filename = u'doc/X9348K-X-02 互提资料单 送电.docx'

    print filename
    # filename = r'd:/test1.doc'
    if '.docx' in filename:
        docx2txt(filename)
        newfilename = filename.replace(u'.docx', u'.txt')
    elif '.doc' in filename:
        doc2txt(filename)
        newfilename = filename.replace(u'.doc', u'.txt')
    read_file = open(newfilename,'r')
    text = read_file.read()
    text = text.replace('  ',' ')
    text = text.replace(' ','')
    text = text.decode('utf-8')
    # seg_list = jieba.cut(text, cut_all=False)
    # print("Full Mode: " + "/ ".join(seg_list))  # 全模式
    seg_list = jieba.lcut(text, cut_all=False) #jieba分成了词
    tok = nltk.word_tokenize(text)  #nltk分成了句子
    # print seg_list
    print seg_list
    print tok
    #print seg_list.count(u'电缆')

